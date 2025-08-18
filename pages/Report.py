# pages/Report.py
# Report Builder — generate a DOCX synthesis from the consolidated session dataframe
# - Uses only local data (st.session_state['consolidated_df'] or 'consolidated')
# - Preview sections in-app before download
# - Noise-aware cleaning (removes TOC/indices/roman numerals/dotted leaders)
# - Extractive summaries (TF-IDF centroid by default; SBERT optional if available)
# - Cited evidence snippets by year/region/filename
# - Graceful fallbacks (no red stack traces; tips instead)

from __future__ import annotations
import io
import re
import unicodedata
from typing import List, Tuple, Dict

import numpy as np
import pandas as pd
import streamlit as st

# Optional (faster/better summaries if available)
HAVE_SBERT = False
try:
    from sentence_transformers import SentenceTransformer
    HAVE_SBERT = True
except Exception:
    pass

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.preprocessing import normalize

# DOCX
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------- Session helpers ----------------
def get_consolidated_df() -> pd.DataFrame | None:
    df1 = st.session_state.get("consolidated_df", None)
    if isinstance(df1, pd.DataFrame) and not df1.empty:
        return df1
    df2 = st.session_state.get("consolidated", None)
    if isinstance(df2, pd.DataFrame) and not df2.empty:
        return df2
    return None

# ---------------- Text utils ----------------
ROMAN_RE = re.compile(r"\b(?=[MDCLXVI])M{0,4}(CM|CD|D?C{0,3})(XC|XL|L?X{0,3})(IX|IV|V?I{0,3})\b", re.I)
DOT_LEADERS = re.compile(r"\.{2,}")                           # dotted leaders: "......"
NUM_COL_RE = re.compile(r"\b\d{1,4}\b")                       # bare numbers
IDX_LINE_RE = re.compile(r"^\s*(\d+(\.\d+)*|[IVXLC]+)\s+[^\s].*$", re.I)  # "3.1 Something", "IV Something"

HEADER_PAT = re.compile(
    r"\b(table of contents|contents|toc|annex|appendix|chapter\s+\d+|figure|fig\.)\b",
    re.I,
)

def fold_accents(s: str) -> str:
    if s is None:
        return ""
    s = s.replace("\r", " ").replace("\t", " ")
    s = (s.replace("\u2010","-").replace("\u2011","-").replace("\u2012","-")
           .replace("\u2013","-").replace("\u2014","-").replace("\u2015","-"))
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"\s+", " ", s).strip()
    return s

def clean_for_summary(text: str) -> str:
    """Remove TOC/indices/roman numerals/dotted leaders/standalone numbers."""
    if not isinstance(text, str):
        return ""
    t = text.replace("\xa0", " ")
    # kill dotted leaders and page-number-ish crud
    t = DOT_LEADERS.sub(" ", t)
    # drop classic header tokens
    t = HEADER_PAT.sub(" ", t)
    # remove standalone roman numerals and lone numbers (keep digits in words)
    t = ROMAN_RE.sub(" ", t)
    t = NUM_COL_RE.sub(lambda m: " " if len(m.group(0)) <= 4 else m.group(0), t)
    # drop index-like lines entirely
    lines = []
    for line in t.splitlines():
        if IDX_LINE_RE.match(line.strip()):
            continue
        lines.append(line)
    t = "\n".join(lines)
    # fold accents and collapse whitespace
    t = fold_accents(t)
    t = re.sub(r"\s+", " ", t).strip()
    return t

def sent_split(text: str) -> List[str]:
    text = re.sub(r"\s+", " ", str(text))
    sents = re.split(r"(?<=[\.\!\?])\s+", text)
    out = [s.strip() for s in sents if s and len(s.strip()) > 2]
    return out

# ---------------- Summarization ----------------
@st.cache_resource(show_spinner=False)
def _load_sbert():
    if not HAVE_SBERT:
        return None
    try:
        return SentenceTransformer("all-MiniLM-L6-v2")
    except Exception:
        return None

def rank_sentences_tfidf(sents: List[str], topk: int = 8) -> List[Tuple[float, str]]:
    """Centroid-based extractive summary with TF-IDF (robust on CPU)."""
    if not sents:
        return []
    # Vectorize with safe defaults (avoid min_df/max_df errors)
    vec = TfidfVectorizer(strip_accents="unicode", ngram_range=(1,2), min_df=1, max_df=1.0)
    X = vec.fit_transform(sents).astype(np.float32)
    if X.shape[0] == 0 or X.shape[1] == 0:
        return []
    X = normalize(X)
    centroid = np.asarray(X.mean(axis=0)).ravel()
    scores = X @ centroid
    scores = np.asarray(scores).ravel()
    order = np.argsort(-scores)[: min(topk, len(sents))]
    return [(float(scores[i]), sents[i]) for i in order]

def rank_sentences_sbert(sents: List[str], topk: int = 8) -> List[Tuple[float, str]]:
    model = _load_sbert()
    if model is None or not sents:
        return rank_sentences_tfidf(sents, topk=topk)
    emb = model.encode(sents, normalize_embeddings=True, show_progress_bar=False)
    centroid = emb.mean(axis=0)
    centroid = centroid / (np.linalg.norm(centroid) + 1e-12)
    scores = emb @ centroid
    order = np.argsort(-scores)[: min(topk, len(sents))]
    return [(float(scores[i]), sents[i]) for i in order]

def summarize_block(texts: List[str], engine: str = "auto", max_sentences: int = 8) -> List[str]:
    """Summarize a list of documents: split→clean→pool sentences→rank→pick top."""
    joined = " ".join(clean_for_summary(t) for t in texts if str(t).strip())
    sents = sent_split(joined)
    if not sents:
        return []
    if engine.lower().startswith("sbert") and _load_sbert() is not None:
        ranked = rank_sentences_sbert(sents, topk=max_sentences)
    else:
        ranked = rank_sentences_tfidf(sents, topk=max_sentences)
    # Keep original order of selected top sentences for readability
    picked = set(s for _, s in ranked)
    ordered = [s for s in sents if s in picked][:max_sentences]
    return ordered

# ---------------- Evidence snippet selection ----------------
def evidence_snippets(df: pd.DataFrame, text_col: str, per_group: int = 3,
                      group_cols: List[str] = None) -> pd.DataFrame:
    """Pick a few representative sentences per group with simple TF-IDF centroid."""
    group_cols = group_cols or ["region", "year"]
    # normalize meta cols
    tmp = pd.DataFrame({
        "text": df[text_col].fillna("").astype(str),
        "filename": (df["filename"] if "filename" in df.columns else pd.Series(np.arange(len(df)).astype(str))),
        "year": df["year"] if "year" in df.columns else np.nan,
        "region": df["unicef_region"] if "unicef_region" in df.columns else (df["region"] if "region" in df.columns else "Unknown"),
        "country": df["country"] if "country" in df.columns else ""
    })
    rows = []
    for keys, g in tmp.groupby([c for c in group_cols if c in tmp.columns], dropna=False):
        texts = g["text"].tolist()
        sents = []
        srcs = []
        for _, r in g.iterrows():
            for s in sent_split(clean_for_summary(r["text"])):
                sents.append(s)
                srcs.append((r["filename"], r.get("country", ""), r.get("region", ""), r.get("year", "")))
        if not sents:
            continue
        ranked = rank_sentences_tfidf(sents, topk=per_group)
        for score, sent in ranked:
            fn, ctry, reg, yr = srcs[sents.index(sent)]
            rec = {"snippet": sent, "score": round(float(score), 4),
                   "filename": fn, "country": ctry, "region": reg, "year": yr}
            # unpack multi-index keys safely
            if isinstance(keys, tuple):
                for i, col in enumerate([c for c in group_cols if c in tmp.columns]):
                    rec[col] = keys[i]
            else:
                rec[group_cols[0]] = keys
            rows.append(rec)
    return pd.DataFrame(rows)

# ---------------- DOCX builder ----------------
def build_docx_report(
    title: str,
    exec_summary: List[str],
    methods_notes: List[str],
    key_themes: List[str],
    evidence_df: pd.DataFrame,
    footer_note: str = "Generated with TextLab (local, offline)."
) -> bytes:
    doc = Document()

    # Title page
    h1 = doc.add_heading(title, level=0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Auto-generated report summary")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing

    # (Optional) TOC placeholder for Word to populate after opening:
    doc.add_paragraph("Table of Contents (update in Word: References → Update Table)")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    if exec_summary:
        for s in exec_summary:
            doc.add_paragraph(s)
    else:
        doc.add_paragraph("No summary could be generated (insufficient text after cleaning).")

    # Methods
    doc.add_heading("2. Methods (Brief)", level=1)
    for m in methods_notes:
        doc.add_paragraph(f"• {m}")

    # Key Themes (lightweight)
    doc.add_heading("3. Key Themes (Extractive)", level=1)
    if key_themes:
        for t in key_themes:
            doc.add_paragraph(f"• {t}")
    else:
        doc.add_paragraph("No salient themes detected.")

    # Evidence by Region/Year
    doc.add_heading("4. Evidence by Region/Year", level=1)
    if not evidence_df.empty:
        # group and list
        grp_cols = [c for c in ["region", "year"] if c in evidence_df.columns]
        for keys, g in evidence_df.groupby(grp_cols, dropna=False):
            title = " — ".join([f"{col}: {val}" for col, val in zip(grp_cols, keys if isinstance(keys, tuple) else [keys])])
            doc.add_heading(title or "Group", level=2)
            for _, r in g.iterrows():
                bullet = doc.add_paragraph(style=None)
                run = bullet.add_run(f"{r['snippet']}")
                bullet.add_run(f"  (Source: {r.get('filename','')}, {r.get('country','')}, {r.get('region','')}, {r.get('year','')})").italic = True
    else:
        doc.add_paragraph("No evidence snippets available.")

    # Footer
    doc.add_page_break()
    doc.add_paragraph(footer_note).alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ---------------- Streamlit UI ----------------
st.set_page_config(page_title="Report Builder", page_icon="📝", layout="wide")
st.title("📝 Report Builder (from session)")

df = get_consolidated_df()
if df is None:
    st.error("No consolidated dataframe in session. Run the main page first.")
    st.stop()

with st.expander("Preview consolidated dataframe", expanded=False):
    st.dataframe(df.head(20), use_container_width=True)

# Column selectors
cols = list(df.columns)
text_defaults = [c for c in ["Findings of the evaluation", "full_text", "text"] if c in cols]
text_col = st.selectbox(
    "Primary text column for summaries",
    options=cols,
    index=cols.index(text_defaults[0]) if text_defaults else 0
)

# Filters
left, right = st.columns(2)
with left:
    # try to infer meta columns
    year_col = "year" if "year" in cols else None
    region_col = "unicef_region" if "unicef_region" in cols else ("region" if "region" in cols else None)
    country_col = "country" if "country" in cols else None

    # Simple filter widgets (optional)
    years = sorted(pd.to_numeric(df.get(year_col, pd.Series([])), errors="coerce").dropna().unique().tolist()) if year_col else []
    year_filter = st.multiselect("Filter by year (optional)", options=years, default=years)
with right:
    regions = sorted(df.get(region_col, pd.Series([])).dropna().unique().tolist()) if region_col else []
    region_filter = st.multiselect("Filter by region (optional)", options=regions, default=regions)

# Apply filters
work = df.copy()
if year_col and year_filter:
    work = work[pd.to_numeric(work[year_col], errors="coerce").isin(year_filter)]
if region_col and region_filter:
    work = work[work[region_col].isin(region_filter)]

if work.empty:
    st.info("No rows match your filters. Clear filters or pick a different column.")
    st.stop()

# Summary settings
st.markdown("---")
c1, c2, c3 = st.columns(3)
with c1:
    max_exec_sents = st.slider("Executive summary length (sentences)", 4, 20, 10, 1)
with c2:
    snippets_per_group = st.slider("Evidence snippets per Region/Year group", 1, 8, 3, 1)
with c3:
    engine = st.radio("Summary engine", ["Auto (SBERT→TF-IDF)", "TF-IDF only"], index=0)

# Compute previews (quietly guard everything)
try:
    # Executive summary from the selected text column
    exec_summary = summarize_block(work[text_col].astype(str).tolist(),
                                   engine="sbert" if engine.startswith("Auto") else "tfidf",
                                   max_sentences=max_exec_sents)

    # Quick "themes": reuse the top summary sentences as bullets
    key_themes = exec_summary[: min(8, len(exec_summary))]

    # Evidence table: top sentences per Region/Year
    group_cols = [c for c in ["region", "year"] if c in work.columns or c in ["unicef_region"]]
    # normalize meta columns for the evidence function
    if "region" not in work.columns and "unicef_region" in work.columns:
        work = work.rename(columns={"unicef_region": "region"})
    evidence = evidence_snippets(work.assign(text=work[text_col]), "text",
                                 per_group=snippets_per_group,
                                 group_cols=["region", "year"])
except Exception:
    exec_summary, key_themes = [], []
    evidence = pd.DataFrame()
    st.info("Tip: Not enough clean text for summaries or snippets. Try a richer column (e.g., full_text).")

# Preview panels
st.subheader("Preview — Executive Summary")
if exec_summary:
    for i, s in enumerate(exec_summary, 1):
        st.markdown(f"{i}. {s}")
else:
    st.caption("No summary available.")

st.subheader("Preview — Evidence Snippets (cited)")
if not evidence.empty:
    st.dataframe(evidence[["region","year","filename","snippet","score","country"]],
                 use_container_width=True, height=360)
else:
    st.caption("No evidence snippets available.")

# Build DOCX
st.markdown("---")
report_title = st.text_input("Report title", "Meta-synthesis Report (Auto)")
methods_notes = [
    "All analysis executed locally in TextLab.",
    "Text cleaned to remove TOC, indices, numerals, and dotted leaders.",
    "Extractive summaries use TF-IDF centroid (SBERT if available) without external APIs.",
    "Evidence snippets ranked by similarity to corpus centroid within groups.",
]
footer = "Generated by TextLab — private, local analysis."

docx_bytes = build_docx_report(
    title=report_title,
    exec_summary=exec_summary,
    methods_notes=methods_notes,
    key_themes=key_themes,
    evidence_df=evidence if not evidence.empty else pd.DataFrame(),
    footer_note=footer
)

st.download_button(
    "⬇️ Download DOCX report",
    data=docx_bytes,
    file_name="TextLab_Report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True
)

st.caption("Open the DOCX in Word and choose **References → Update Table** if you want a full Table of Contents.")
